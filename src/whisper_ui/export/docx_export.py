from __future__ import annotations

import io
import logging

from whisper_ui.core.exceptions import ExportError
from whisper_ui.core.models import TranscriptResult
from whisper_ui.ui.labels import EXPORT_DOCX_HEADING

logger = logging.getLogger(__name__)


class DocxExporter:
    @property
    def format_name(self) -> str:
        return "DOCX"

    @property
    def file_extension(self) -> str:
        return ".docx"

    @property
    def mime_type(self) -> str:
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def export(self, result: TranscriptResult) -> bytes:
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError as err:
            raise ExportError("python-docx is not installed. Install it with: pip install python-docx") from err

        doc = Document()
        doc.add_heading(EXPORT_DOCX_HEADING, level=1)

        current_speaker: str | None = None
        for seg in result.segments:
            if seg.speaker and seg.speaker != current_speaker:
                current_speaker = seg.speaker
                p = doc.add_paragraph()
                run = p.add_run(f"[{current_speaker}]")
                run.bold = True
                run.font.size = Pt(11)

            p = doc.add_paragraph(seg.text)
            p.style.font.size = Pt(10)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
