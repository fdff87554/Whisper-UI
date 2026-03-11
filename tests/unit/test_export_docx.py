from __future__ import annotations

import io

from docx import Document

from whisper_ui.core.models import Segment, TranscriptResult
from whisper_ui.export.docx_export import DocxExporter


def test_docx_export_basic():
    result = TranscriptResult(
        segments=[
            Segment(start=0.0, end=1.0, text="Hello", speaker="SPEAKER_00"),
            Segment(start=1.0, end=2.0, text="World", speaker="SPEAKER_01"),
        ],
        language="en",
        duration=2.0,
    )
    exporter = DocxExporter()
    data = exporter.export(result)
    assert isinstance(data, bytes)
    assert len(data) > 0

    doc = Document(io.BytesIO(data))
    texts = [p.text for p in doc.paragraphs]
    assert any("Hello" in t for t in texts)
    assert any("World" in t for t in texts)
    assert any("[SPEAKER_00]" in t for t in texts)
    assert any("[SPEAKER_01]" in t for t in texts)


def test_docx_export_no_speakers():
    result = TranscriptResult(
        segments=[Segment(start=0.0, end=1.0, text="No speaker")],
    )
    exporter = DocxExporter()
    data = exporter.export(result)
    doc = Document(io.BytesIO(data))
    texts = [p.text for p in doc.paragraphs]
    assert any("No speaker" in t for t in texts)
    assert not any("SPEAKER" in t for t in texts)


def test_docx_export_empty():
    result = TranscriptResult(segments=[])
    exporter = DocxExporter()
    data = exporter.export(result)
    assert isinstance(data, bytes)
    assert len(data) > 0


def test_docx_format_metadata():
    exporter = DocxExporter()
    assert exporter.format_name == "DOCX"
    assert exporter.file_extension == ".docx"
    assert "wordprocessingml" in exporter.mime_type
